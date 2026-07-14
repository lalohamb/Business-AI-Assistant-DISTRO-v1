#!/usr/bin/env python3
"""Index Obsidian vault and system/client files into PostgreSQL + pgvector."""

import os
import logging
import psycopg2
from pathlib import Path
from dotenv import load_dotenv

env_path = Path(__file__).resolve().parent.parent / ".env"
load_dotenv(env_path)

# Log to both console and file
log_dir = Path(__file__).resolve().parent.parent / "logs"
log_dir.mkdir(exist_ok=True)
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S",
    handlers=[
        logging.StreamHandler(),
        logging.FileHandler(log_dir / "index_vault.log", mode="a"),
    ],
)
log = logging.getLogger(__name__)

BASE_PATH = os.getenv("BASE_PATH")
ACTIVE_CLIENT = os.getenv("ACTIVE_CLIENT", "demo-company")
EMBEDDING_PROVIDER = os.getenv("EMBEDDING_PROVIDER", "ollama")
EMBEDDING_MODEL = os.getenv("EMBEDDING_MODEL", "nomic-embed-text")
EMBEDDING_DIMENSIONS = int(os.getenv("EMBEDDING_DIMENSIONS", "768"))
OLLAMA_BASE_URL = os.getenv("OLLAMA_BASE_URL", "http://localhost:11434")

EXCLUDE_DIRS = {"admin", "logs", "backups", "docker", "postgres", "node_modules", ".git", "venv"}
EXCLUDE_EXTENSIONS = {".key", ".pem"}
EXCLUDE_FILES = {".env"}

INDEX_PATHS = [
    os.path.join(BASE_PATH, "system"),
    os.path.join(BASE_PATH, "clients", ACTIVE_CLIENT),
]

DB_CONFIG = {
    "host": os.getenv("PG_HOST", "localhost"),
    "port": int(os.getenv("PG_PORT", "5432")),
    "user": os.getenv("PG_USER", "admin"),
    "password": os.getenv("PG_PASSWORD", "strongpassword"),
    "dbname": os.getenv("PG_DATABASE", "businessassistant"),
}


SUPPORTED_EXTENSIONS = {".md", ".txt", ".pdf", ".docx", ".xlsx", ".csv", ".html", ".eml"}


def extract_text(filepath):
    """Extract text content from supported file formats."""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in (".md", ".txt", ".csv", ".eml"):
        with open(filepath, "r", errors="ignore") as f:
            return f.read().strip()
    elif ext == ".pdf":
        import fitz
        doc = fitz.open(filepath)
        return "\n".join(page.get_text() for page in doc).strip()
    elif ext == ".docx":
        from docx import Document
        doc = Document(filepath)
        return "\n".join(p.text for p in doc.paragraphs).strip()
    elif ext == ".xlsx":
        from openpyxl import load_workbook
        wb = load_workbook(filepath, read_only=True, data_only=True)
        lines = []
        for ws in wb.worksheets:
            for row in ws.iter_rows(values_only=True):
                lines.append(" | ".join(str(c) if c is not None else "" for c in row))
        return "\n".join(lines).strip()
    elif ext == ".html":
        from bs4 import BeautifulSoup
        with open(filepath, "r", errors="ignore") as f:
            return BeautifulSoup(f.read(), "html.parser").get_text(separator="\n").strip()
    return ""


def get_files(paths):
    """Collect all indexable files, excluding admin/logs/backups/docker/postgres/.git."""
    files = []
    for base in paths:
        if not os.path.exists(base):
            continue
        for root, dirs, filenames in os.walk(base):
            dirs[:] = [d for d in dirs if d not in EXCLUDE_DIRS]
            for f in filenames:
                if f in EXCLUDE_FILES:
                    continue
                ext = os.path.splitext(f)[1].lower()
                if ext in EXCLUDE_EXTENSIONS:
                    continue
                if ext in SUPPORTED_EXTENSIONS:
                    files.append(os.path.join(root, f))
    return files


def chunk_text(text, chunk_size=512, overlap=64):
    """Split text into chunks, preferring markdown section boundaries."""
    import re
    # Split on markdown section separators first
    sections = re.split(r'\n---\n|\n## ', text)
    chunks = []
    for i, section in enumerate(sections):
        # Re-add the heading prefix stripped by split (except first)
        if i > 0 and not section.startswith('#'):
            section = '## ' + section
        section = section.strip()
        if not section:
            continue
        # If section fits in one chunk, keep it whole
        if len(section) <= chunk_size:
            chunks.append(section)
        else:
            # Fall back to size-based splitting for large sections
            start = 0
            while start < len(section):
                end = start + chunk_size
                chunks.append(section[start:end])
                start += chunk_size - overlap
    return [c for c in chunks if c.strip()]


def get_embedding(text):
    """Get embedding vector from configured provider."""
    if EMBEDDING_PROVIDER == "ollama":
        import requests
        resp = requests.post(
            f"{OLLAMA_BASE_URL}/api/embeddings",
            json={"model": EMBEDDING_MODEL, "prompt": text},
        )
        resp.raise_for_status()
        return resp.json()["embedding"]
    else:
        raise NotImplementedError(f"Embedding provider \"{EMBEDDING_PROVIDER}\" not yet supported.")


def ensure_schema(cur):
    """Ensure tables exist with correct embedding dimensions. Recreate if dimensions changed."""
    cur.execute("SELECT EXISTS(SELECT 1 FROM information_schema.tables WHERE table_name='rag_chunks')")
    table_exists = cur.fetchone()[0]

    if table_exists:
        # Check current dimension
        cur.execute("""
            SELECT atttypmod FROM pg_attribute
            WHERE attrelid = 'rag_chunks'::regclass AND attname = 'embedding'
        """)
        row = cur.fetchone()
        current_dim = row[0] if row else 0
        if current_dim != EMBEDDING_DIMENSIONS:
            log.info(f"Embedding dimensions changed ({current_dim} → {EMBEDDING_DIMENSIONS}). Recreating tables...")
            cur.execute("DROP TABLE IF EXISTS rag_chunks CASCADE")
            cur.execute("DROP TABLE IF EXISTS rag_documents CASCADE")
            table_exists = False

    if not table_exists:
        cur.execute("CREATE EXTENSION IF NOT EXISTS vector")
        cur.execute("""
            CREATE TABLE IF NOT EXISTS rag_documents (
                id SERIAL PRIMARY KEY,
                client_name VARCHAR(255) NOT NULL,
                source_path TEXT NOT NULL,
                title VARCHAR(500),
                created_at TIMESTAMP DEFAULT NOW()
            )
        """)
        cur.execute(f"""
            CREATE TABLE IF NOT EXISTS rag_chunks (
                id SERIAL PRIMARY KEY,
                document_id INTEGER REFERENCES rag_documents(id) ON DELETE CASCADE,
                client_name VARCHAR(255) NOT NULL,
                source_path TEXT NOT NULL,
                title VARCHAR(500),
                chunk_text TEXT NOT NULL,
                embedding vector({EMBEDDING_DIMENSIONS}),
                created_at TIMESTAMP DEFAULT NOW()
            )
        """)
        cur.execute("CREATE INDEX IF NOT EXISTS idx_chunks_client ON rag_chunks(client_name)")
        log.info(f"Created tables with embedding dimension={EMBEDDING_DIMENSIONS}")


def index():
    """Main indexing pipeline."""
    files = get_files(INDEX_PATHS)
    log.info(f"Found {len(files)} files to index.")
    log.info(f"Embedding model: {EMBEDDING_MODEL} (dimensions={EMBEDDING_DIMENSIONS})")

    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()

    ensure_schema(cur)
    conn.commit()

    cur.execute("DELETE FROM rag_chunks WHERE client_name = %s", (ACTIVE_CLIENT,))
    cur.execute("DELETE FROM rag_documents WHERE client_name = %s", (ACTIVE_CLIENT,))

    for filepath in files:
        try:
            content = extract_text(filepath)
        except Exception as e:
            log.warning(f"  Skipping {filepath}: {e}")
            continue

        if not content:
            continue

        title = os.path.basename(filepath)
        rel_path = os.path.relpath(filepath, BASE_PATH)

        cur.execute(
            "INSERT INTO rag_documents (client_name, source_path, title) VALUES (%s, %s, %s) RETURNING id",
            (ACTIVE_CLIENT, rel_path, title),
        )
        doc_id = cur.fetchone()[0]

        chunks = chunk_text(content)
        for chunk in chunks:
            # Prepend source context to improve embedding relevance
            embed_text = f"[Source: {title} for {ACTIVE_CLIENT}] {chunk}"
            try:
                embedding = get_embedding(embed_text)
            except Exception as e:
                log.warning(f"  Embedding failed for chunk in {title}: {e}")
                continue

            cur.execute(
                "INSERT INTO rag_chunks (document_id, client_name, source_path, title, chunk_text, embedding) VALUES (%s, %s, %s, %s, %s, %s)",
                (doc_id, ACTIVE_CLIENT, rel_path, title, chunk, embedding),
            )

        log.info(f"  Indexed: {rel_path} ({len(chunks)} chunks)")

    conn.commit()

    # Rebuild ivfflat index after data is inserted
    cur.execute("SELECT COUNT(*) FROM rag_chunks")
    row_count = cur.fetchone()[0]
    lists = max(1, min(int(row_count ** 0.5), row_count // 10))
    log.info(f"Rebuilding ivfflat index (lists={lists} for {row_count} rows)...")
    cur.execute("DROP INDEX IF EXISTS idx_chunks_embedding")
    cur.execute(f"CREATE INDEX idx_chunks_embedding ON rag_chunks USING ivfflat (embedding vector_cosine_ops) WITH (lists = {lists})")
    conn.commit()

    cur.close()
    conn.close()
    log.info("Indexing complete.")


if __name__ == "__main__":
    index()
